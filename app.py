import os, io, re, shutil, tempfile, secrets, base64, hashlib, datetime
from flask import Flask, render_template, render_template_string, request, session, redirect, url_for, send_file
from dotenv import load_dotenv
from jinja2 import Environment, FileSystemLoader, TemplateNotFound
import pdfkit  # HTML→PDF via wkhtmltopdf if present; Playwright fallback inside helper
# PDF compose
from reportlab.pdfgen import canvas as rl_canvas
from reportlab.lib.utils import ImageReader
from PyPDF2 import PdfReader, PdfWriter
import pdfplumber  # (anchor utils optional)
# DOCX → HTML (for Contract)
from docx import Document
import mammoth
# Data
import pandas as pd
import numpy as np
import requests
from flask import url_for
from playwright.sync_api import sync_playwright
import tempfile, os
# -------------------------------------------------------------------
# Load env + Flask
# -------------------------------------------------------------------
load_dotenv()
app = Flask(__name__)
app.secret_key = os.getenv("FLASK_SECRET_KEY") or secrets.token_hex(16)
app.config["MAX_CONTENT_LENGTH"] = 200 * 1024 * 1024
app.config["PROPAGATE_EXCEPTIONS"] = True

# -------------------------------------------------------------------
# Notion basics (you already use this token elsewhere)
# -------------------------------------------------------------------
NOTION_TOKEN = os.getenv("NOTION_TOKEN")
NOTION_API = "https://api.notion.com/v1"
NOTION_VERSION = "2022-06-28"

def notion_headers(json=True):
    if not NOTION_TOKEN:
        raise RuntimeError("NOTION_TOKEN not set in environment.")
    h = {
        "Authorization": f"Bearer {NOTION_TOKEN}",
        "Notion-Version": NOTION_VERSION,
    }
    if json:
        h["Content-Type"] = "application/json"
    return h

# -------------------------------------------------------------------
# Your Notion DBs (Pages 1–5 content lives elsewhere; this file uses the token + DB ids)
# If you also use a helper NotionClient in a separate module, that's fine; this app works standalone.
# -------------------------------------------------------------------
NOTION_DATABASE_ID = os.getenv("NOTION_DATABASE_ID")  # Page 1 DB id (Borrower) — used by other routes

NOTION_DATABASE_ID_PAGE2 = os.getenv("NOTION_DATABASE_ID_PAGE2") or NOTION_DATABASE_ID
NOTION_DATABASE_ID_PAGE3 = os.getenv("NOTION_DATABASE_ID_PAGE3") or NOTION_DATABASE_ID_PAGE2
NOTION_DATABASE_ID_PAGE4 = os.getenv("NOTION_DATABASE_ID_PAGE4") or NOTION_DATABASE_ID_PAGE3  # Loan Application
NOTION_DATABASE_ID_PAGE5 = os.getenv("NOTION_DATABASE_ID_PAGE5") or NOTION_DATABASE_ID_PAGE4

# (If you previously imported a NotionClient class, you can keep using it.
# For clarity here, we do minimal direct POST/GETs with `requests`.)

# -------------------------------------------------------------------
# Paths / Constants
# -------------------------------------------------------------------
ROOT_DIR     = os.path.abspath(os.path.dirname(__file__))
TEMPLATE_DIR = ROOT_DIR
STATIC_DIR   = os.path.join(ROOT_DIR, "static")
CONTRACTS_DIR= os.path.join(ROOT_DIR, "contracts")

# KFS (HTML template) + upload property
KFS_TEMPLATE_FILENAME = os.getenv("KFS_TEMPLATE", "template.html")
KFS_FILES_PROP_NAME   = os.getenv("KFS_FILES_PROP_NAME", "Key Fact Sheet")

# KFS fixed-point signature location (pt; 1 pt = 1/72 inch)
KFS_SIG_PAGE_INDEX = int(os.getenv("KFS_SIG_PAGE_INDEX", "1"))     # default to page 2
KFS_SIG_X_PT       = float(os.getenv("KFS_SIG_X_PT", "0"))
KFS_SIG_Y_PT       = float(os.getenv("KFS_SIG_Y_PT", "400"))
KFS_SIG_WIDTH_PT   = float(os.getenv("KFS_SIG_WIDTH_PT", "220"))

# Contract template + upload property
CONTRACT_TEMPLATE_PATH   = os.path.join(CONTRACTS_DIR, "Microfinancing Loan Agreement.docx")
CONTRACT_FILES_PROP_NAME = os.getenv("CONTRACT_FILES_PROP_NAME", "Executed Contract")

# Contract fixed-point signature location (independent from KFS)
CONTRACT_SIG_PAGE_INDEX = int(os.getenv("CONTRACT_SIG_PAGE_INDEX", "0"))
CONTRACT_SIG_X_PT       = float(os.getenv("CONTRACT_SIG_X_PT", "420"))
CONTRACT_SIG_Y_PT       = float(os.getenv("CONTRACT_SIG_Y_PT", "100"))
CONTRACT_SIG_WIDTH_PT   = float(os.getenv("CONTRACT_SIG_WIDTH_PT", "220"))

# -------------------------------------------------------------------
# Notion DBs for Ledger Export (Disbursements / Collections)
# -------------------------------------------------------------------
NOTION_DB_DISBURSEMENTS = os.getenv("NOTION_DB_DISBURSEMENTS")  # required
NOTION_DB_COLLECTIONS   = os.getenv("NOTION_DB_COLLECTIONS")    # required
# Column/property names (exact)
PROP_DISB_DATE   = os.getenv("PROP_DISB_DATE",   "Date")
PROP_DISB_AMOUNT = os.getenv("PROP_DISB_AMOUNT", "Amount")
PROP_COLL_DATE   = os.getenv("PROP_COLL_DATE",   "Date")
PROP_COLL_AMOUNT = os.getenv("PROP_COLL_AMOUNT", "Amount")
# Starting corpus default
STARTING_CORPUS  = float(os.getenv("STARTING_CORPUS", "7500000"))

# -------------------------------------------------------------------
# Minimal Notion Query Helper
# -------------------------------------------------------------------
def query_notion_db(db_id: str, filter_block=None, sorts=None, page_size=100):
    payload = {"page_size": page_size}
    if filter_block:
        payload["filter"] = filter_block
    if sorts:
        payload["sorts"] = sorts
    results, cursor = [], None
    while True:
        if cursor:
            payload["start_cursor"] = cursor
        r = requests.post(f"{NOTION_API}/databases/{db_id}/query",
                          headers=notion_headers(), json=payload, timeout=30)
        if r.status_code >= 300:
            raise RuntimeError(f"Notion query failed: {r.status_code} {r.text}")
        data = r.json()
        results.extend(data.get("results", []))
        if not data.get("has_more"):
            break
        cursor = data.get("next_cursor")
    return results

def extract_date(prop: dict):
    if not isinstance(prop, dict): return None
    # date
    if "date" in prop and prop["date"] and prop["date"].get("start"):
        return pd.to_datetime(prop["date"]["start"]).date()
    # formula(date)
    if "formula" in prop and prop["formula"].get("type") == "date":
        d = prop["formula"]["date"]
        if d and d.get("start"):
            return pd.to_datetime(d["start"]).date()
    # rollup(array of dates)
    if "rollup" in prop:
        r = prop["rollup"]
        if r.get("type") == "array" and r["array"]:
            item = r["array"][0]
            if "date" in item and item["date"]:
                return pd.to_datetime(item["date"]["start"]).date()
    # title/rich_text fallback
    for key in ("title","rich_text"):
        if key in prop and prop[key]:
            try:
                return pd.to_datetime(prop[key][0].get("plain_text")).date()
            except Exception:
                return None
    return None

def extract_number(prop: dict):
    if not isinstance(prop, dict): return 0.0
    if "number" in prop and prop["number"] is not None:
        return float(prop["number"])
    if "formula" in prop and prop["formula"].get("type") == "number":
        n = prop["formula"]["number"]
        return 0.0 if n is None else float(n)
    if "rollup" in prop and prop["rollup"].get("type") == "number":
        n = prop["rollup"]["number"]
        return 0.0 if n is None else float(n)
    for key in ("title","rich_text"):
        if key in prop and prop[key]:
            txt = prop[key][0].get("plain_text","").replace(",","")
            try:
                return float(txt)
            except Exception:
                pass
    return 0.0

def fetch_df_from_notion(db_id: str, date_prop: str, amt_prop: str, date_from=None, date_to=None) -> pd.DataFrame:
    flt = None
    if date_from or date_to:
        cond = {}
        if date_from: cond["on_or_after"] = pd.to_datetime(date_from).date().isoformat()
        if date_to:   cond["on_or_before"] = pd.to_datetime(date_to).date().isoformat()
        flt = {"property": date_prop, "date": cond}
    rows = query_notion_db(db_id, filter_block=flt, sorts=[{"property": date_prop, "direction": "ascending"}])
    out = []
    for pg in rows:
        props = pg.get("properties", {})
        d = extract_date(props.get(date_prop, {}))
        a = extract_number(props.get(amt_prop, {}))
        if d is not None:
            out.append({"Date": d, "Amount": a})
    if not out:
        return pd.DataFrame(columns=["Date","Amount"])
    return pd.DataFrame(out).groupby("Date", as_index=False)["Amount"].sum()

def build_ledger(disb_df: pd.DataFrame, coll_df: pd.DataFrame, start_corpus: float,
                 rng_start=None, rng_end=None) -> pd.DataFrame:
    """
    Minimal daily ledger for display/export:
      Date, Deployed, Collected, Net Inflow (Day), Corpus Remaining

    Net Inflow (Day) = Collected - Deployed  (positive when collections exceed deployments)
    """
    import datetime
    # Nothing to compute and no user range → return empty shaped frame
    if disb_df.empty and coll_df.empty and not rng_start and not rng_end:
        return pd.DataFrame(columns=[
            "Date", "Deployed", "Collected", "Net Inflow (Day)", "Corpus Remaining"
        ])

    # ---- safe min/max dates ----
    mins = []
    if not disb_df.empty and pd.notna(disb_df["Date"].min()):
        mins.append(pd.to_datetime(disb_df["Date"].min()).date())
    if not coll_df.empty and pd.notna(coll_df["Date"].min()):
        mins.append(pd.to_datetime(coll_df["Date"].min()).date())
    if rng_start:
        mins.append(pd.to_datetime(rng_start).date())
    min_date = mins[0] if mins else datetime.date.today()
    for d in mins:
        if d < min_date:
            min_date = d

    maxs = [pd.to_datetime(rng_end).date()] if rng_end else [datetime.date.today()]
    if not disb_df.empty and pd.notna(disb_df["Date"].max()):
        maxs.append(pd.to_datetime(disb_df["Date"].max()).date())
    if not coll_df.empty and pd.notna(coll_df["Date"].max()):
        maxs.append(pd.to_datetime(coll_df["Date"].max()).date())
    max_date = maxs[0]
    for d in maxs:
        if d > max_date:
            max_date = d
    if max_date < min_date:
        max_date = min_date

    # ---- daily calendar + sums ----
    calendar = pd.DataFrame({"Date": pd.date_range(min_date, max_date, freq="D").date})
    disb_daily = (disb_df.rename(columns={"Amount": "Deployed"}) if not disb_df.empty
                  else pd.DataFrame(columns=["Date", "Deployed"]))
    coll_daily = (coll_df.rename(columns={"Amount": "Collected"}) if not coll_df.empty
                  else pd.DataFrame(columns=["Date", "Collected"]))

    tr = calendar.merge(disb_daily, on="Date", how="left").merge(coll_daily, on="Date", how="left")
    tr[["Deployed", "Collected"]] = tr[["Deployed", "Collected"]].fillna(0.0)

    # ---- metrics (keep only requested ones) ----
    tr["Net Inflow (Day)"] = tr["Collected"] - tr["Deployed"]
    # internal cumulatives (not shown) needed to compute corpus
    tr["_cum_disbursed"]  = tr["Deployed"].cumsum()
    tr["_cum_collected"]  = tr["Collected"].cumsum()
    tr["_outstanding"]    = tr["_cum_disbursed"] - tr["_cum_collected"]
    tr["Corpus Remaining"] = start_corpus - tr["_outstanding"]

    # final view
    cols = ["Date", "Deployed", "Collected", "Net Inflow (Day)", "Corpus Remaining"]
    return tr[cols]



def excel_bytes_from_df(df: pd.DataFrame, start_corpus: float, rng_start, rng_end) -> bytes:
    bio = io.BytesIO()
    cfg = pd.DataFrame({
        "Key":   ["Starting Corpus","From","To","Generated At (UTC)"],
        "Value": [start_corpus,
                  (rng_start.isoformat() if isinstance(rng_start, datetime.date) else (rng_start or "")),
                  (rng_end.isoformat()   if isinstance(rng_end,   datetime.date) else (rng_end   or "")),
                  datetime.datetime.utcnow().isoformat(timespec="seconds")+"Z"]
    })
    out = df.copy()
    for c in out.columns:
        if c != "Date":
            out[c] = out[c].round(2)
    with pd.ExcelWriter(bio, engine="xlsxwriter", datetime_format="yyyy-mm-dd", date_format="yyyy-mm-dd") as writer:
        cfg.to_excel(writer, sheet_name="Config", index=False)
        out.to_excel(writer, sheet_name="Daily Tracker", index=False)
    bio.seek(0)
    return bio.getvalue()

# -------------------------------------------------------------------
# HTML→PDF helper (wkhtmltopdf or Playwright)
# -------------------------------------------------------------------
def resolve_wkhtmltopdf_path() -> str | None:
    candidates = [
        os.getenv("WKHTMLTOPDF_CMD"),
        shutil.which("wkhtmltopdf"),
        "/opt/homebrew/bin/wkhtmltopdf",
        "/usr/local/bin/wkhtmltopdf",
        "/usr/bin/wkhtmltopdf",
    ]
    for c in candidates:
        if c and os.path.exists(c):
            return c
    return None

def html_to_pdf_bytes(html_str: str) -> bytes:
    # Always use Playwright Chromium for PDF rendering
    with tempfile.NamedTemporaryFile(
        mode="w", suffix=".html", dir=os.getcwd(), delete=False, encoding="utf-8"
    ) as f:
        f.write(html_str)
        tmp_html = f.name

    pdf_bytes = b""
    with sync_playwright() as p:
        browser = p.chromium.launch(args=["--no-sandbox"])  # --no-sandbox helps in Render/Docker
        page = browser.new_page()
        page.goto("file://" + tmp_html, wait_until="domcontentloaded", timeout=60000)
        pdf_bytes = page.pdf(format="A4", print_background=True)
        browser.close()

    try:
        os.remove(tmp_html)
    except Exception:
        pass

    return pdf_bytes


# -------------------------------------------------------------------
# KFS rendering (HTML template + logo)
# -------------------------------------------------------------------
def render_kfs_html(data: dict, extra: dict | None = None) -> str:
    env = Environment(loader=FileSystemLoader(TEMPLATE_DIR), autoescape=False)
    try:
        tpl = env.get_template(KFS_TEMPLATE_FILENAME)
    except TemplateNotFound:
        raise RuntimeError(f"Cannot find KFS template '{KFS_TEMPLATE_FILENAME}' in {TEMPLATE_DIR}")
    ctx = dict(data)
    if extra: ctx.update(extra)
    return tpl.render(ctx)

def parse_field(prop):
    if not isinstance(prop, dict): return ""
    if "title" in prop:     return prop["title"][0]["plain_text"] if prop["title"] else ""
    if "rich_text" in prop: return prop["rich_text"][0]["plain_text"] if prop["rich_text"] else ""
    if "select" in prop and prop["select"]: return prop["select"]["name"]
    if "date" in prop and prop["date"]:     return prop["date"]["start"]
    if "number" in prop:    return "" if prop["number"] is None else str(prop["number"])
    if "formula" in prop:
        f = prop["formula"]
        if f["type"] == "string": return f["string"] or ""
        if f["type"] == "number": return "" if f["number"] is None else str(f["number"])
        if f["type"] == "date":   return f["date"]["start"] if f["date"] else ""
    if "rollup" in prop:
        r = prop["rollup"]
        if r["type"] == "array" and r["array"]:
            first = r["array"][0]
            for key in ("title","rich_text","text"):
                if key in first and first[key]:
                    return first[key][0].get("plain_text","")
            if "number" in first and first["number"] is not None: return str(first["number"])
            if "date" in first and first["date"]: return first["date"]["start"]
            return ""
        if r["type"] == "number": return "" if r["number"] is None else str(r["number"])
    if "relation" in prop and prop["relation"]:
        return prop["relation"][0].get("id","")
    if "unique_id" in prop:
        pfx = prop["unique_id"].get("prefix") or ""; num = prop["unique_id"].get("number") or ""
        return f"{pfx}{num}"
    return ""

def fetch_loan_page_values(loan_page_id: str) -> dict:
    # Get the page (Loan Application DB assumed in NOTION_DATABASE_ID_PAGE4)
    r = requests.get(f"{NOTION_API}/pages/{loan_page_id}", headers=notion_headers(), timeout=30)
    if r.status_code >= 300:
        raise RuntimeError(f"Get page failed: {r.status_code} {r.text}")
    p = r.json().get("properties", {})
    g = lambda k: parse_field(p.get(k, {}))
    return {
        "Loan_Application_ID": g("Loan Application ID"),
        "Borrower_Name": g("Full Name"),
        "Co_Borrower_Name": g("Co-borrower"),
        "Loan_Type": g("Loan Type"),
        "Sanction_Date": g("Sanction Date"),
        "Amount_Sanctioned": g("Amount Sanctioned"),
        "Tenure": g("Tenure (Months)"),
        "Interest_Rate": g("Interest Rate (Yearly)"),
        "EMI_Frequency": g("Repayment Frequency") or g("Frequency"),
        "EMI_Amount": g("EMI Amount"),
        "Total_Repayable_Amount": g("Outstanding Amount "),
        "First_EMI_Date": g("Start Date"),
        "Last_EMI_Date": g("End Date"),
        "Processing_Fee": g("Processing_Fee"),
        "Insurance_Fee": g("Insurance_Fee"),
        "Stamp_Duty": g("Stamp_Duty"),
        "Foreclosure_Charges": g("Foreclosure_Clauses"),
        "Disbursement_Date": g("Disbursement Date"),
        "Bank_Account_Details": g("Bank_Account_Details"),
        "Mode": g("Mode"),
        "Mandate_Status": g("Mandate_Status"),
        "Credit_Officer_Name": g("Credit Officer Assigned"),
    }

# -------------------------------------------------------------------
# Signature stamping helpers (fixed-point)
# -------------------------------------------------------------------
def decode_data_url_png(data_url: str) -> bytes:
    if not data_url or "," not in data_url: raise RuntimeError("Invalid signature data.")
    return base64.b64decode(data_url.split(",", 1)[1])

def stamp_signature_at_point(pdf_bytes: bytes, signature_png: bytes, attn_text: str,
                             page_index: int = 0, x_pt: float = 420, y_pt: float = 100, width_pt: float = 220) -> bytes:
    """Place signature at exact (x,y) in PDF points (origin bottom-left)."""
    reader = PdfReader(io.BytesIO(pdf_bytes))
    if page_index >= len(reader.pages):
        # clamp to last page
        page_index = max(0, len(reader.pages) - 1)
    page = reader.pages[page_index]
    pw, ph = float(page.mediabox.width), float(page.mediabox.height)
    sig_w = width_pt; sig_h = sig_w * 0.33
    x = max(6, min(x_pt, pw - sig_w - 6))
    y = max(6, min(y_pt, ph - sig_h - 6))
    overlay_stream = io.BytesIO()
    c = rl_canvas.Canvas(overlay_stream, pagesize=(pw, ph))
    img = ImageReader(io.BytesIO(signature_png))
    c.drawImage(img, x, y, width=sig_w, height=sig_h, mask='auto')
    c.setFont("Helvetica", 9); c.setFillGray(0.2)
    c.drawString(x, y - 12, attn_text[:140])
    c.showPage(); c.save()
    overlay_pdf = PdfReader(io.BytesIO(overlay_stream.getvalue()))
    page.merge_page(overlay_pdf.pages[0])
    out = PdfWriter()
    for p in reader.pages: out.add_page(p)
    out_buf = io.BytesIO(); out.write(out_buf)
    return out_buf.getvalue()

# -------------------------------------------------------------------
# Small shared helpers (forms)
# -------------------------------------------------------------------
def form_fields(fields_def):
    out = []
    for f in fields_def:
        f2 = dict(f); f2["id"] = f["notion_prop"]; out.append(f2)
    return out

def validate(req, fields):
    errors, clean = {}, {}
    for f in fields:
        fid = f["id"]; ftype = f["type"]; required = f.get("required", False)
        if ftype == "multi_select":
            raw = req.form.getlist(fid)
        elif ftype == "files":
            raw = req.files.getlist(fid)
        elif ftype == "relation_session":
            raw = None
        else:
            raw = req.form.get(fid)
        if required and ftype != "relation_session":
            missing = (
                (ftype == "multi_select" and len(raw) == 0) or
                (ftype == "files" and all((not fs or not getattr(fs, "filename", "")) for fs in (raw or []))) or
                (ftype not in ["multi_select", "files"] and (raw is None or str(raw).strip() == ""))
            )
            if missing:
                errors[fid] = "This field is required."
                clean[fid] = raw
                continue
        clean[fid] = raw
    return clean, errors

# -------------------------------------------------------------------
# PAGES 1–4 (Borrower, Field Verification, Risk Remarks, Loan Application)
# (Keep your existing schema.py etc. — here we show only the route skeletons you were using)
# -------------------------------------------------------------------
from schema import FIELDS
from schema_page2 import FIELDS_PAGE2
from schema_page3 import FIELDS_PAGE3
from schema_page4 import FIELDS_PAGE4
from schema_page5 import FIELDS_PAGE5

@app.route("/", methods=["GET","POST"])
def form():
    fields = form_fields(FIELDS)
    errors, success_id = {}, None
    try:
        if request.method == "POST":
            clean, errors = validate(request, fields)
            if not errors:
                # create page in NOTION_DATABASE_ID (Borrower)
                props = {}
                for f in fields:
                    v = clean.get(f["id"])
                    # map quickly: number/text/date/select/checkbox/url/email/phone/status
                    # Minimal title mapping:
                    t = f["type"]; name = f["notion_prop"]
                    if t == "title":
                        props[name] = {"title": [{"type": "text", "text": {"content": (v or "")}}]}
                    elif t == "number":
                        props[name] = {"number": float(v) if v not in (None,"") else None}
                    elif t == "date":
                        props[name] = {"date": {"start": v}} if v else {"date": None}
                    elif t == "select":
                        props[name] = {"select": {"name": v}} if v else {"select": None}
                    elif t == "multi_select":
                        props[name] = {"multi_select": [{"name": x} for x in (v or [])]}
                    elif t == "checkbox":
                        props[name] = {"checkbox": (v == "on" or v is True)}
                    elif t == "url":
                        props[name] = {"url": v or None}
                    elif t == "email":
                        props[name] = {"email": v or None}
                    elif t == "phone":
                        props[name] = {"phone_number": v or None}
                    elif t == "rich_text":
                        props[name] = {"rich_text": [{"type":"text","text":{"content": v or ""}}]}
                    elif t == "files":
                        files_arr = []
                        for fs in (v or []):
                            if fs and getattr(fs, "filename", ""):
                                # Example: save to temp and upload to Notion
                                init = requests.post(f"{NOTION_API}/file_uploads",
                                                     headers=notion_headers(),
                                                     json={"filename": fs.filename, "mode": "single_part"})
                                init.raise_for_status()
                                up_id = init.json()["id"]
                                send = requests.post(f"{NOTION_API}/file_uploads/{up_id}/send",
                                                     headers=notion_headers(json=False),
                                                     files={"file": (fs.filename, fs, fs.mimetype or "application/octet-stream")})
                                send.raise_for_status()
                                files_arr.append({
                                    "name": fs.filename,
                                    "type": "file_upload",
                                    "file_upload": {"id": up_id}
                                })
                        props[name] = {"files": files_arr}
                    else:
                        props[name] = {"rich_text": [{"type":"text","text":{"content": str(v) if v is not None else ""}}]}
                r = requests.post(f"{NOTION_API}/pages", headers=notion_headers(), json={
                    "parent": {"database_id": NOTION_DATABASE_ID},
                    "properties": props
                }, timeout=30)
                if r.status_code >= 300:
                    errors["__all__"] = f"Notion error: {r.status_code} {r.text}"
                else:
                    success_id = r.json().get("id")
                    session["notion_page_id"] = success_id
                    # save Borrower Name title for later display
                    for f in fields:
                        if f.get("type") == "title" and f.get("name") == "Borrower Name":
                            session["borrower_name"] = request.form.get(f["notion_prop"], "")
                            break
                    return redirect(url_for("form_page2"))
    except Exception as e:
        errors["__all__"] = f"{type(e).__name__}: {e}"
    return render_template("form.html", title="Borrower Information", fields=fields,
                           errors=errors, success=False, success_id=success_id,
                           csrf_token=secrets.token_hex(16))

@app.route("/page2", methods=["GET","POST"])
def form_page2():
    fields = form_fields(FIELDS_PAGE2)
    errors, success_id = {}, None
    linked_id = session.get("notion_page_id"); borrower_name = session.get("borrower_name")
    try:
        if request.method == "POST":
            clean, errors = validate(request, fields)
            if not errors:
                props = {}
                for f in fields:
                    fid, ftype = f["id"], f["type"]; name = f["notion_prop"]
                    if ftype == "relation_session":
                        if not linked_id:
                            errors["__all__"] = "Missing borrower link. Submit Page 1 first."
                            break
                        props[name] = {"relation": [{"id": linked_id}]}
                    else:
                        v = clean.get(fid)
                        if ftype == "date":
                            props[name] = {"date": {"start": v}} if v else {"date": None}

                        elif ftype == "datetime":
                            props[name] = {"date": {"start": v}} if v else {"date": None}

                        elif ftype == "number":
                            props[name] = {"number": float(v) if v not in (None,"") else None}

                        elif ftype == "select":
                            props[name] = {"select": {"name": v}} if v else {"select": None}

                        elif ftype == "checkbox":
                            props[name] = {"checkbox": (v == "on" or v is True)}

                        elif ftype == "url":
                            props[name] = {"url": v if v else None}

                        elif ftype == "rich_text":
                            props[name] = {"rich_text": [{"type":"text","text":{"content": v or ""}}]}

                        else:
                            # Fallback for unexpected types
                            props[name] = {"rich_text": [{"type":"text","text":{"content": str(v) if v is not None else ""}}]}

                if not errors:
                    r = requests.post(f"{NOTION_API}/pages", headers=notion_headers(), json={
                        "parent": {"database_id": NOTION_DATABASE_ID_PAGE2}, "properties": props
                    }, timeout=30)
                    if r.status_code >= 300:
                        errors["__all__"] = f"Notion error: {r.status_code} {r.text}"
                    else:
                        success_id = r.json().get("id")
                        session["verification_page_id"] = success_id
                        return redirect(url_for("form_page3"))
    except Exception as e:
        errors["__all__"] = f"{type(e).__name__}: {e}"
    return render_template("form.html", title="Field Verification", subtitle="Step 2",
                           fields=fields, errors=errors, success=False, success_id=success_id,
                           csrf_token=secrets.token_hex(16), linked_id=linked_id, borrower_name=borrower_name)

@app.route("/page3", methods=["GET","POST"])
def form_page3():
    fields = form_fields(FIELDS_PAGE3)
    errors, success_id = {}, None
    linked_id = session.get("notion_page_id"); borrower_name = session.get("borrower_name")
    try:
        if request.method == "POST":
            clean, errors = validate(request, fields)
            if not errors:
                props = {}
                for f in fields:
                    fid, ftype = f["id"], f["type"]; name = f["notion_prop"]
                    if ftype == "relation_session":
                        if not linked_id:
                            errors["__all__"] = "Missing borrower link. Submit Page 1 first."
                            break
                        props[name] = {"relation": [{"id": linked_id}]}
                    else:
                        v = clean.get(fid)
                        if ftype == "rich_text":
                            props[name] = {"rich_text":[{"type":"text","text":{"content": v or ""}}]}
                        elif ftype == "select":
                            props[name] = {"select":{"name": v}} if v else {"select": None}
                        else:
                            props[name] = {"rich_text":[{"type":"text","text":{"content": str(v) if v is not None else ""}}]}
                if not errors:
                    r = requests.post(f"{NOTION_API}/pages", headers=notion_headers(), json={
                        "parent":{"database_id": NOTION_DATABASE_ID_PAGE3}, "properties": props
                    }, timeout=30)
                    if r.status_code >= 300:
                        errors["__all__"] = f"Notion error: {r.status_code} {r.text}"
                    else:
                        success_id = r.json().get("id")
                        session["risk_remarks_page_id"] = success_id
                        return redirect(url_for("form_page4"))
    except Exception as e:
        errors["__all__"] = f"{type(e).__name__}: {e}"
    return render_template("form.html", title="Risk Remarks", subtitle="Step 3",
                           fields=fields, errors=errors, success=False, success_id=success_id,
                           csrf_token=secrets.token_hex(16), linked_id=linked_id, borrower_name=borrower_name)

@app.route("/page4", methods=["GET","POST"])
def form_page4():
    fields = form_fields(FIELDS_PAGE4)
    errors, success_id = {}, None
    linked_id = session.get("notion_page_id"); borrower_name = session.get("borrower_name")
    try:
        if request.method == "POST":
            clean, errors = validate(request, fields)
            if not errors:
                props = {}
                for f in fields:
                    fid, ftype = f["id"], f["type"]; name = f["notion_prop"]
                    if ftype == "relation_session":
                        if not linked_id:
                            errors["__all__"] = "Missing borrower link. Submit Page 1 first."
                            break
                        props[name] = {"relation":[{"id": linked_id}]}

                    elif ftype in ("date", "datetime"):
                        v = clean.get(fid)
                        props[name] = {"date": {"start": v}} if v else {"date": None}

                    elif ftype == "number":
                        v = clean.get(fid)
                        props[name] = {"number": float(v) if v not in (None,"") else None}

                    elif ftype == "select":
                        v = clean.get(fid)
                        props[name] = {"select": {"name": v}} if v else {"select": None}

                    elif ftype == "checkbox":
                        v = clean.get(fid)
                        props[name] = {"checkbox": (v == "on" or v is True)}

                    elif ftype == "url":
                        v = clean.get(fid)
                        props[name] = {"url": v if v else None}

                    elif ftype == "rich_text":
                        v = clean.get(fid)
                        props[name] = {"rich_text":[{"type":"text","text":{"content": v or ""}}]}

                    else:
                        v = clean.get(fid)
                        props[name] = {"rich_text":[{"type":"text","text":{"content": str(v) if v is not None else ""}}]}

                if not errors:
                    r = requests.post(f"{NOTION_API}/pages", headers=notion_headers(), json={
                        "parent":{"database_id": NOTION_DATABASE_ID_PAGE4}, "properties": props
                    }, timeout=30)
                    if r.status_code >= 300:
                        errors["__all__"] = f"Notion error: {r.status_code} {r.text}"
                    else:
                        success_id = r.json().get("id")
                        session["loan_application_page_id"] = success_id
                        return redirect(url_for("kfs_sign"))
    except Exception as e:
        errors["__all__"] = f"{type(e).__name__}: {e}"
    return render_template("form.html", title="Loan Application Internal", subtitle="Step 4",
                           fields=fields, errors=errors, success=False, success_id=success_id,
                           csrf_token=secrets.token_hex(16), linked_id=linked_id, borrower_name=borrower_name)

# -------------------------------------------------------------------
# KFS preview + submit (fixed-point signing)
# -------------------------------------------------------------------
@app.route("/kfs-sign", methods=["GET"])
def kfs_sign():
    loan_id = session.get("loan_application_page_id")
    if not loan_id: return redirect(url_for("form_page4"))
    try:
        data = fetch_loan_page_values(loan_id)
        logo_http = url_for("static", filename="ub-portfolio-logo.png")
        kfs_html = render_kfs_html(data, extra={"logo_src": logo_http})
    except Exception as e:
        session["kfs_error"] = f"KFS render failed: {e}"
        return redirect(url_for("form_page5"))
    sign_ctx = {
        "ip": request.headers.get("X-Forwarded-For", request.remote_addr or ""),
        "now_utc": datetime.datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S"),
    }
    # Use your white sign_kfs.html (already in your project)
    return render_template("sign_kfs.html", title="Review & Sign KFS",
                           kfs_html=kfs_html, borrower_name=session.get("borrower_name"),
                           sign_ctx=sign_ctx)

@app.route("/kfs-sign/submit", methods=["POST"])
def kfs_sign_submit():
    loan_id = session.get("loan_application_page_id")
    if not loan_id: return redirect(url_for("form_page4"))
    sig_data_url = request.form.get("sig_data_url","")
    signer_name  = request.form.get("signer_name","").strip()
    signer_email = request.form.get("signer_email","").strip()
    attn_ip      = request.form.get("attn_ip","")
    attn_ts_utc  = request.form.get("attn_ts_utc","")
    if not (sig_data_url and signer_name and signer_email):
        session["kfs_error"] = "Signature, name, and email are required."
        return redirect(url_for("kfs_sign"))
    try:
        data = fetch_loan_page_values(loan_id)
        logo_http = url_for("static", filename="ub-portfolio-logo.png", _external=True)
        html = render_kfs_html(data, extra={"logo_src": logo_http})
        print("Start KFS sign submit")
        pdf = html_to_pdf_bytes(html)
        print("PDF generated")
        sig_png = decode_data_url_png(sig_data_url)
        print("Signature decoded")
        attn_text = f"Signed by {signer_name} <{signer_email}> at {attn_ts_utc} UTC · IP {attn_ip}"
        # KFS: fixed point (can be different from contract)
        signed_pdf = stamp_signature_at_point(pdf, sig_png, attn_text,
                                              page_index=KFS_SIG_PAGE_INDEX,
                                              x_pt=KFS_SIG_X_PT, y_pt=KFS_SIG_Y_PT,
                                              width_pt=KFS_SIG_WIDTH_PT)
        print("Signature stamped")
        pdf_sha256 = hashlib.sha256(signed_pdf).hexdigest()
        fname = f"signed_kfs_{data.get('Loan_Application_ID') or loan_id}.pdf"
        # Upload file
        # 1) Create file upload
        init = requests.post(f"{NOTION_API}/file_uploads", headers=notion_headers(), json={
            "filename": fname, "mode":"single_part"
        }, timeout=30)
        init.raise_for_status()
        up_id = init.json()["id"]
        # 2) Send bytes
        send = requests.post(f"{NOTION_API}/file_uploads/{up_id}/send",
                             headers=notion_headers(json=False),
                             files={"file": (fname, io.BytesIO(signed_pdf), "application/pdf")},
                             timeout=180)
        send.raise_for_status()
        # 3) Patch page property
        patch = requests.patch(f"{NOTION_API}/pages/{loan_id}",
                               headers=notion_headers(), json={
            "properties": {
                KFS_FILES_PROP_NAME: {
                    "files": [{"name": fname, "type":"file_upload", "file_upload":{"id": up_id}}]
                },
                "Signed By Name": {"rich_text":[{"type":"text","text":{"content": signer_name}}]},
                "Signed By Email": {"email": signer_email},
                "Signed IP": {"rich_text":[{"type":"text","text":{"content": attn_ip}}]},
                "Signed At": {"date":{"start": attn_ts_utc}},
                "KFS SHA256": {"rich_text":[{"type":"text","text":{"content": pdf_sha256}}]},
                "KFS Signed": {"checkbox": True}
            }
        }, timeout=30)
        patch.raise_for_status()
    except Exception as e:
        session["kfs_error"] = f"KFS signing/upload failed: {e}"
    return redirect(url_for("form_page5"))

# -------------------------------------------------------------------
# Page 5 — after submit, you may redirect to Contract Sign (Page 6) or finish.
# -------------------------------------------------------------------
@app.route("/page5", methods=["GET","POST"])
def form_page5():
    fields = form_fields(FIELDS_PAGE5)
    errors, success = {}, False
    kfs_err = session.pop("kfs_error", None)
    if kfs_err: errors["__all__"] = kfs_err
    borrower_id = session.get("notion_page_id"); borrower_name = session.get("borrower_name")
    try:
        if request.method == "POST":
            props = {}
            for f in fields:
                sid = session.get(f.get("session_key",""))
                if not sid:
                    errors["__all__"] = f"Missing link for {f['name']}. Please complete prior steps."
                    break
                props[f["notion_prop"]] = {"relation":[{"id": sid}]}
            if not errors:
                # create page in NOTION_DATABASE_ID_PAGE5
                r = requests.post(f"{NOTION_API}/pages", headers=notion_headers(), json={
                    "parent":{"database_id": NOTION_DATABASE_ID_PAGE5},
                    "properties": props
                }, timeout=30)
                if r.status_code >= 300:
                    errors["__all__"] = f"Notion error: {r.status_code} {r.text}"
                else:
                    return redirect(url_for("contract_sign"))  # go to contract
    except Exception as e:
        errors["__all__"] = f"{type(e).__name__}: {e}"
    return render_template("form.html", title="Income Assessment Form", subtitle="Step 5",
                           fields=fields, errors=errors, success=success, success_id=None,
                           csrf_token=secrets.token_hex(16), linked_id=borrower_id, borrower_name=borrower_name)

# -------------------------------------------------------------------
# Page 6 — Contract Sign + Thank You
# -------------------------------------------------------------------
@app.route("/contract-sign", methods=["GET"])
def contract_sign():
    try:
        # Fill your DOCX placeholders -> HTML preview
        loan_id = session.get("loan_application_page_id")
        if not loan_id:
            return redirect(url_for("form_page4"))
        # Build mapping
        data = fetch_loan_page_values(loan_id)
        borrower_name = session.get("borrower_name","")
        now = datetime.datetime.utcnow()
        mapping = {
            "Borrower’s name": borrower_name,
            "Borrower's name": borrower_name,
            "amount": data.get("Amount_Sanctioned") or "",
            "date": now.strftime("%d"),
            "month": now.strftime("%B"),
            "year": now.strftime("%Y"),
        }
        if not os.path.exists(CONTRACT_TEMPLATE_PATH):
            raise RuntimeError(f"Contract template not found at {CONTRACT_TEMPLATE_PATH}")
        # Replace <<key>> in DOCX and convert to HTML
        doc = Document(CONTRACT_TEMPLATE_PATH)
        def rep_runs(runs, needle, repl):
            for r in runs:
                if needle in r.text:
                    r.text = r.text.replace(needle, repl)
        for p in doc.paragraphs:
            for k,v in mapping.items(): rep_runs(p.runs, f"<<{k}>>", str(v))
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for k,v in mapping.items(): rep_runs(p.runs, f"<<{k}>>", str(v))
        tmpdoc = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        doc.save(tmpdoc.name)
        with open(tmpdoc.name, "rb") as f:
            contract_html = mammoth.convert_to_html(f).value
        try: os.remove(tmpdoc.name)
        except Exception: pass
    except Exception as e:
        session["kfs_error"] = f"Contract render failed: {e}"
        return redirect(url_for("form_page5"))
    sign_ctx = {
        "ip": request.headers.get("X-Forwarded-For", request.remote_addr or ""),
        "now_utc": datetime.datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S"),
    }
    return render_template("contract_sign.html", title="Review & Sign Loan Agreement",
                           contract_html=contract_html, borrower_name=session.get("borrower_name"),
                           sign_ctx=sign_ctx)

@app.route("/contract-sign/submit", methods=["POST"])
def contract_sign_submit():
    loan_id = session.get("loan_application_page_id")
    if not loan_id: return redirect(url_for("form_page4"))
    sig_data_url = request.form.get("sig_data_url","")
    signer_name  = request.form.get("signer_name","").strip()
    signer_email = request.form.get("signer_email","").strip()
    attn_ip      = request.form.get("attn_ip","")
    attn_ts_utc  = request.form.get("attn_ts_utc","")
    if not (sig_data_url and signer_name and signer_email):
        session["kfs_error"] = "Signature, name, and email are required."
        return redirect(url_for("contract_sign"))
    try:
        # Re-render filled contract HTML → PDF
        borrower_name = session.get("borrower_name","")
        data = fetch_loan_page_values(loan_id)
        now = datetime.datetime.utcnow()
        mapping = {
            "Borrower’s name": borrower_name,
            "Borrower's name": borrower_name,
            "amount": data.get("Amount_Sanctioned") or "",
            "date": now.strftime("%d"),
            "month": now.strftime("%B"),
            "year": now.strftime("%Y"),
        }
        doc = Document(CONTRACT_TEMPLATE_PATH)
        def rep_runs(runs, needle, repl):
            for r in runs:
                if needle in r.text:
                    r.text = r.text.replace(needle, repl)
        for p in doc.paragraphs:
            for k,v in mapping.items(): rep_runs(p.runs, f"<<{k}>>", str(v))
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for k,v in mapping.items(): rep_runs(p.runs, f"<<{k}>>", str(v))
        tmpdoc = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        doc.save(tmpdoc.name)
        with open(tmpdoc.name, "rb") as f:
            contract_html = mammoth.convert_to_html(f).value
        try: os.remove(tmpdoc.name)
        except Exception: pass

        pdf  = html_to_pdf_bytes(contract_html)
        sig_png   = decode_data_url_png(sig_data_url)
        attn_text = f"Signed by {signer_name} <{signer_email}> at {attn_ts_utc} UTC · IP {attn_ip}"
        signed_pdf = stamp_signature_at_point(pdf, sig_png, attn_text,
                                              page_index=CONTRACT_SIG_PAGE_INDEX,
                                              x_pt=CONTRACT_SIG_X_PT, y_pt=CONTRACT_SIG_Y_PT,
                                              width_pt=CONTRACT_SIG_WIDTH_PT)
        # Upload to Loan Application page: Executed Contract
        fname = f"signed_loan_agreement_{loan_id}.pdf"
        init = requests.post(f"{NOTION_API}/file_uploads", headers=notion_headers(), json={
            "filename": fname, "mode":"single_part"
        }, timeout=30)
        init.raise_for_status()
        up_id = init.json()["id"]
        send = requests.post(f"{NOTION_API}/file_uploads/{up_id}/send",
                             headers=notion_headers(json=False),
                             files={"file": (fname, io.BytesIO(signed_pdf), "application/pdf")},
                             timeout=180)
        send.raise_for_status()
        patch = requests.patch(f"{NOTION_API}/pages/{loan_id}",
                               headers=notion_headers(), json={
            "properties": {
                CONTRACT_FILES_PROP_NAME: {
                    "files": [{"name": fname, "type":"file_upload", "file_upload":{"id": up_id}}]
                }
            }
        }, timeout=30)
        patch.raise_for_status()
    except Exception as e:
        session["kfs_error"] = f"Contract signing/upload failed: {e}"
        return redirect(url_for("form_page5"))
    # Thank You
    return redirect(url_for("thank_you"))

@app.route("/thank-you", methods=["GET"])
def thank_you():
    return render_template("thank_you.html",
                           title="Thank you!",
                           borrower_name=session.get("borrower_name"),
                           loan_id=session.get("loan_application_page_id"))

# -------------------------------------------------------------------
# Notion-Powered Ledger Page + Excel download
# -------------------------------------------------------------------

@app.route("/export-ledger", methods=["GET"])
def export_ledger():
    """
    Render the daily ledger as an HTML table (full-width, full-height).
    Add ?download=1 to get Excel instead.
      ?start=YYYY-MM-DD&end=YYYY-MM-DD&start_corpus=7500000
    """
    if not (NOTION_DB_DISBURSEMENTS and NOTION_DB_COLLECTIONS):
        return "Please set NOTION_DB_DISBURSEMENTS and NOTION_DB_COLLECTIONS in .env", 500

    q_start = request.args.get("start")
    q_end   = request.args.get("end")
    q_download = request.args.get("download")
    try:
        start_corpus = float(request.args.get("start_corpus") or STARTING_CORPUS)
    except Exception:
        start_corpus = STARTING_CORPUS

    # Fetch Notion data
    try:
        disb_df = fetch_df_from_notion(NOTION_DB_DISBURSEMENTS, PROP_DISB_DATE, PROP_DISB_AMOUNT, q_start, q_end)
        coll_df = fetch_df_from_notion(NOTION_DB_COLLECTIONS,   PROP_COLL_DATE, PROP_COLL_AMOUNT, q_start, q_end)
    except Exception as e:
        return f"Failed to query Notion: {e}", 500

    rng_start = pd.to_datetime(q_start).date() if q_start else None
    rng_end   = pd.to_datetime(q_end).date() if q_end else None
    ledger = build_ledger(disb_df, coll_df, start_corpus, rng_start, rng_end)

    # --- Summary metrics ---
    if not ledger.empty:
        last_corpus_remaining = float(ledger["Corpus Remaining"].iloc[-1])
        amount_actually_deployed = float(start_corpus - last_corpus_remaining)
        total_deployed = float(ledger["Deployed"].sum())
        total_collected = float(ledger["Collected"].sum())
        returns_realised = ((total_deployed + total_collected) / amount_actually_deployed) if amount_actually_deployed > 0 else None
    else:
        last_corpus_remaining = start_corpus
        amount_actually_deployed = 0.0
        total_deployed = 0.0
        total_collected = 0.0
        returns_realised = None

    # Download Excel?
    if q_download:
        xbytes = excel_bytes_from_df(ledger, start_corpus, rng_start or "", rng_end or "")
        fname = f"Daily_Corpus_Tracker_{(rng_start or (ledger['Date'].min() if not ledger.empty else 'NA'))}_{(rng_end or datetime.date.today())}.xlsx"
        return send_file(io.BytesIO(xbytes), as_attachment=True,
                         download_name=str(fname).replace(":","-"),
                         mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

    # View model (rounded numbers)
    view = ledger.copy()
    for c in view.columns:
        if c != "Date":
            view[c] = view[c].round(2)

    meta = {
        "start": (rng_start.isoformat() if rng_start else ""),
        "end":   (rng_end.isoformat()   if rng_end   else ""),
        "start_corpus": start_corpus,
        "row_count": len(view),
        "last_updated": datetime.datetime.utcnow().isoformat(timespec="seconds")+"Z"
    }
    summary = {
        "amount_actually_deployed": amount_actually_deployed,
        "returns_realised": returns_realised,
        "last_corpus_remaining": last_corpus_remaining,
        "total_deployed": total_deployed,
        "total_collected": total_collected,
    }

    # Full-viewport layout + sticky toolbar, and a Summary card at the bottom
    html = """
    {% extends "base.html" %}
    {% block content %}

      <style>
        html, body { height: 100%; }
        .wrap { max-width: 100% !important; padding-left: 24px; padding-right: 24px; }
        .ledger-shell { display: flex; flex-direction: column; height: calc(100vh - 24px); gap: 12px; }
        .ledger-header .nf-title { font-size: 28px; }
        .ledger-header .nf-sub   { margin-top: 6px; }

        .toolbar { position: sticky; top: 0; z-index: 5; backdrop-filter: blur(4px); -webkit-backdrop-filter: blur(4px); }
        .filters { display: grid; grid-template-columns: repeat(6, minmax(160px, 1fr)); gap: 12px; align-items: end; }
        @media (max-width: 1100px) { .filters { grid-template-columns: repeat(2, minmax(160px, 1fr)); } }

        .table-wrap { flex: 1; min-height: 0; overflow: auto; border: 1px solid var(--border); border-radius: 12px; background: var(--card, #0f172a); }
        table.ledger { width: 100%; border-collapse: collapse; table-layout: fixed; }
        table.ledger thead th { position: sticky; top: 0; z-index: 2; padding: 10px 12px; text-align: left; border-bottom: 1px solid var(--border); background: var(--card, #0f172a); white-space: nowrap; }
        table.ledger tbody td { padding: 8px 12px; border-bottom: 1px solid var(--border); white-space: nowrap; }
        table.ledger tbody td.num { text-align: right; font-variant-numeric: tabular-nums; }

        .summary { display: grid; grid-template-columns: repeat(2, minmax(240px, 1fr)); gap: 12px; }
        @media (max-width: 820px) { .summary { grid-template-columns: 1fr; } }
        .summary .card { background: var(--card, #0f172a); border: 1px solid var(--border); border-radius: 12px; padding: 16px; }
        .summary h3 { margin: 0 0 8px 0; font-size: 16px; }
        .kv { display:flex; justify-content: space-between; padding: 6px 0; border-bottom: 1px dashed var(--border); }
        .kv:last-child { border-bottom: none; }
        .kv .k { color: var(--muted); }
        .kv .v { font-variant-numeric: tabular-nums; }
      </style>

      <div class="ledger-shell">

        <section class="nf-header ledger-header">
          <h1 class="nf-title">Daily Corpus Tracker</h1>
          <p class="nf-sub">Built from Notion · Last updated: {{ meta.last_updated }}</p>
        </section>

        <div class="nf-field toolbar">
          <form id="filterForm" method="get" action="{{ url_for('export_ledger') }}" class="filters">
            <div>
              <label class="nf-label">From</label>
              <input class="nf-input" type="date" name="start" id="start" value="{{ meta.start }}">
            </div>
            <div>
              <label class="nf-label">To</label>
              <input class="nf-input" type="date" name="end" id="end" value="{{ meta.end }}">
            </div>
            <div>
              <label class="nf-label">Starting Corpus</label>
              <input class="nf-input" type="number" name="start_corpus" id="start_corpus" step="any" value="{{ meta.start_corpus }}">
            </div>
            <div>
              <label class="nf-label">&nbsp;</label>
              <button class="nf-btn" type="submit">Refresh</button>
            </div>
            <div>
              <label class="nf-label">&nbsp;</label>
              <button class="nf-btn" type="button" id="downloadBtn">Download Excel</button>
            </div>
          </form>
        </div>

        <div class="table-wrap">
          {% if table is not none and table.shape[0] > 0 %}
            <table class="ledger">
              <thead>
                <tr>
                  {% for col in table.columns %}
                    <th>{{ col }}</th>
                  {% endfor %}
                </tr>
              </thead>
              <tbody>
                {% for _, row in table.iterrows() %}
                  <tr>
                    {% for col in table.columns %}
                      {% if col == "Date" %}
                        <td>{{ row[col] }}</td>
                      {% else %}
                        <td class="num">{{ "{:,.2f}".format(row[col]) if row[col] == row[col] else "" }}</td>
                      {% endif %}
                    {% endfor %}
                  </tr>
                {% endfor %}
              </tbody>
            </table>
          {% else %}
            <div class="nf-help" style="padding: 12px;">No rows found for the selected range.</div>
          {% endif %}
        </div>

        <div class="nf-help" style="margin-top:6px;">
          Rows: {{ meta.row_count }} · Starting Corpus: {{ "{:,.2f}".format(meta.start_corpus) }}
        </div>

        <!-- Summary card -->
        <div class="summary">
          <div class="card">
            <h3>Summary</h3>
            <div class="kv">
              <div class="k">Amount actually deployed</div>
              <div class="v">
                {% if summary.amount_actually_deployed is not none %}
                  ₹ {{ "{:,.2f}".format(summary.amount_actually_deployed) }}
                {% else %}—{% endif %}
              </div>
            </div>
            <div class="kv">
              <div class="k">Returns Realised</div>
              <div class="v">
                {% if summary.returns_realised is not none %}
                  {{ "{:,.4f}".format(summary.returns_realised) }}
                {% else %}—{% endif %}
              </div>
            </div>
          </div>

          <!-- (Optional) extra card if you want to see components; you can remove this block -->
          <div class="card">
            <h3>Components</h3>
            <div class="kv"><div class="k">Total Deployed</div><div class="v">₹ {{ "{:,.2f}".format(summary.total_deployed) }}</div></div>
            <div class="kv"><div class="k">Total Collected</div><div class="v">₹ {{ "{:,.2f}".format(summary.total_collected) }}</div></div>
            <div class="kv"><div class="k">Corpus Remaining</div><div class="v">₹ {{ "{:,.2f}".format(summary.last_corpus_remaining) }}</div></div>
          </div>
        </div>

      </div>

      <script>
        (function(){
          const btn = document.getElementById('downloadBtn');
          btn.addEventListener('click', function(){
            const start = document.getElementById('start').value;
            const end   = document.getElementById('end').value;
            const sc    = document.getElementById('start_corpus').value;
            const u = new URL(window.location.origin + "{{ url_for('export_ledger') }}");
            if (start) u.searchParams.set('start', start);
            if (end)   u.searchParams.set('end', end);
            if (sc)    u.searchParams.set('start_corpus', sc);
            u.searchParams.set('download', '1');
            window.location = u.toString();
          });
        })();
      </script>

    {% endblock %}
    """
    return render_template_string(html, table=view, meta=meta, summary=summary)

@app.route("/debug/ledger-env")
def debug_ledger_env():
    lines = [
        f"NOTION_DB_DISBURSEMENTS = {NOTION_DB_DISBURSEMENTS}",
        f"NOTION_DB_COLLECTIONS   = {NOTION_DB_COLLECTIONS}",
        f"PROP_DISB_DATE          = {PROP_DISB_DATE}",
        f"PROP_DISB_AMOUNT        = {PROP_DISB_AMOUNT}",
        f"PROP_COLL_DATE          = {PROP_COLL_DATE}",
        f"PROP_COLL_AMOUNT        = {PROP_COLL_AMOUNT}",
        f"STARTING_CORPUS         = {STARTING_CORPUS}",
    ]
    return "<pre>" + "\n".join(lines) + "</pre>"

@app.route("/debug/notion-props/<which>")
def debug_notion_props(which):
    db = NOTION_DB_DISBURSEMENTS if which.lower().startswith("disb") else NOTION_DB_COLLECTIONS
    rows = query_notion_db(db, page_size=1)
    if not rows:
        return "<pre>No rows in that database.</pre>"
    props = rows[0].get("properties", {})
    lines = [f"{k} : {list(v.keys())}" for k, v in props.items()]
    return "<pre>" + "\n".join(lines) + "</pre>"
# -------------------------------------------------------------------
# Links + run
# -------------------------------------------------------------------
if __name__ == "__main__":
    host = "0.0.0.0"  # Always bind to all interfaces
    port = int(os.getenv("PORT", "10000"))  # Render provides PORT
    url = f"http://{host}:{port}/"
    print("\n================= Flask Dev Server =================")
    print(f"→ Page 1 (Borrower): {url}")
    print(f"→ Page 2:            {url}page2")
    print(f"→ Page 3:            {url}page3")
    print(f"→ Page 4:            {url}page4")
    print(f"→ KFS Sign:          {url}kfs-sign")
    print(f"→ Page 5:            {url}page5")
    print(f"→ Contract Sign:     {url}contract-sign")
    print(f"→ Ledger:            {url}export-ledger")
    print("====================================================\n", flush=True)
    app.run(host=host, port=port, debug=False)

